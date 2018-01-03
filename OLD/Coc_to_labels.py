# -*- coding: utf-8 -*-
"""
Created on Fri Sep 09 14:35:31 2016

@author: alex.messina
"""

## Take a generic Excel File of Chain of Custody (CoC) and generate the appropriate labels

## Basic modules
import pandas as pd
import numpy as np

## Document modules
import docx
from docx import *
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import csv

#%%

maindir = "C:/Users/alex.messina/Documents/GitHub/CoC-labels/"


CoC = maindir+'WECK 19L Clean and Blank COC.xls'

form = pd.ExcelFile(CoC).parse(sheetname='STANDARD CHAIN-OF-CUSTODY',header=13,index_col=0,parse_cols='B:F,I:S',skip_footer=15)[1:]
form.columns = ['DATE SAMPLED','TIME SAMPLED','SMPL TYPE','SAMPLE IDENTIFICATION/SITE LOCATION','# OF CONT.','Protocol Cleaning','Ammonia','Nitrate-N','Nitrite-N','Total Kjeldahl Nitrogen','Total phosphorus as P','Total Metals','Mercury','TOC','Cyanide']

form = form[pd.notnull(form.index)].astype(np.str)

#%%
##  Create Document
document = Document(maindir+'template.docx')
#document.add_heading('Labels',0)


df = []
for row in form.iterrows():
    print ''
    print 'SampleID: '+str(row[0])
    print 'DATE SAMPLED: '+row[1]['DATE SAMPLED']+'  TIME SAMPLED: '+row[1]['TIME SAMPLED']

    


    analytes = []
    if row[1]['Ammonia']=='X':
        analytes.append('Ammonia')
    if row[1]['Nitrate-N']=='X':
        analytes.append('Nitrate-N')
    if row[1]['Nitrite-N']=='X':
        analytes.append('Nitrite-N')
    if row[1]['Total Kjeldahl Nitrogen']=='X':
        analytes.append('TKN')
    if row[1]['Total phosphorus as P']=='X':
        analytes.append('TP')
    if row[1]['Total Metals']=='X':
        analytes.append('Total Metals')
    if row[1]['Mercury']=='X':
        analytes.append('Mercury')
    if row[1]['TOC']=='X':
        analytes.append('TOC')
    if row[1]['Cyanide']=='X':
        analytes.append('Cyanide')
        
    df.append(['"SITE ID:  '+str(row[0]) +'","'+ 'DATE SAMPLED:  '+row[1]['DATE SAMPLED']+'    TIME SAMPLED:  '+row[1]['TIME SAMPLED'] +'","'+'Analytes: '+', '.join(analytes) +'"'])  
    
    print 'Analyses: '
    print analytes
    
    document.add_paragraph('SITE ID: '+str(row[0]))
    document.add_paragraph('DATE SAMPLED: '+row[1]['DATE SAMPLED']+'  TIME SAMPLED: '+row[1]['TIME SAMPLED'])
    document.add_paragraph('ANALYSES: '+', '.join(analytes))
    document.add_paragraph('')
    
    
    




document.save(maindir+'labels.docx')
#%%

#csv_out = open('labels.txt','w')
#for item in df:
#    print item[0]
#    if len(item) > 
#    csv_out.write("%s\n"%item[0])
#csv_out.close()


















